from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
BUNDLE = ROOT / "chapter4_final_system_analysis_bundle_2026-05-15"
OUT = BUNDLE / "Section_4_System_Analysis_Final.docx"


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=12)


def configure_page(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.right_margin = Inches(1)
    section.left_margin = Inches(1.5)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    p = normal.paragraph_format
    p.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "ChapterTitle" not in doc.styles:
        style = doc.styles.add_style("ChapterTitle", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(16)
        style.font.bold = True

    if "CustomHeading" not in doc.styles:
        style = doc.styles.add_style("CustomHeading", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if "CaptionStyle" not in doc.styles:
        style = doc.styles.add_style("CaptionStyle", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_after = Pt(4)


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        prefix, rest = bold_prefix, text[len(bold_prefix):]
        r1 = p.add_run(prefix)
        set_run_font(r1, size=12, bold=True)
        r2 = p.add_run(rest)
        set_run_font(r2, size=12)
    else:
        r = p.add_run(text)
        set_run_font(r, size=12)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph(style="CustomHeading")
    r = p.add_run(text)
    set_run_font(r, size=14, bold=True)
    return p


def add_title(doc, text):
    p = doc.add_paragraph(style="ChapterTitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=16, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="CaptionStyle")
    r = p.add_run(text)
    set_run_font(r, size=12, italic=True)
    return p


def add_table(doc, title, headers, rows):
    add_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=12, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=12)
    doc.add_paragraph("")


def add_picture(doc, path, width, caption):
    doc.add_picture(str(path), width=width)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_page(section)
    section.start_type = WD_SECTION.NEW_PAGE

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Page ")
    set_run_font(r, size=12)
    add_page_field(footer, "PAGE")
    r = footer.add_run(" of ")
    set_run_font(r, size=12)
    add_page_field(footer, "NUMPAGES")

    add_title(doc, "CHAPTER 4")
    add_title(doc, "SYSTEM ANALYSIS")

    intro_paragraphs = [
        "System analysis explains how the proposed underwater image enhancement system is organized, what resources it depends on, how data moves through the pipeline, and how the final outputs are produced for research evaluation. In this project, system analysis is especially important because underwater image enhancement is not a single-step correction problem. The input images are affected by wavelength-dependent absorption, scattering, haze, low contrast, and severe color shift. A useful enhancement system must therefore combine reliable data preparation, a stable training setup, and an implementation process that can be repeated and evaluated objectively.",
        "The present work implements a lightweight deep learning framework for underwater image enhancement using a convolutional encoder-decoder with a MobileViT-style bottleneck. The overall system is designed around five core components: dataset description, training setup, hyperparameter specification, software requirements, and implementation process. These components together define how the system is built and why it is suitable for an academic research study. Rather than treating the model as an isolated algorithm, this chapter presents the supporting workflow that makes the model trainable, measurable, and reproducible.",
    ]
    for para in intro_paragraphs:
        add_paragraph(doc, para)

    add_heading(doc, "4.1 Dataset Description")
    dataset_paragraphs = [
        "The performance of an underwater image enhancement system depends strongly on the quality and diversity of the datasets used during training and evaluation. Unlike ordinary land-scene enhancement, underwater image restoration must cope with blue-green dominance, weak red-channel information, turbidity, non-uniform illumination, and reduced visibility caused by suspended particles. For that reason, the proposed system does not rely on a single image source. Instead, it uses benchmark underwater datasets that collectively represent supervised training conditions as well as real-world no-reference evaluation scenarios.",
        "The main training source is the EUVP paired underwater scenes dataset. It is organized into `trainA`, `trainB`, and `validation` folders under the `data/EUVP/Paired/underwater_scenes/` directory. The `trainA` folder stores degraded underwater inputs, while `trainB` contains the corresponding reference targets used for supervised learning. This paired structure is important because the model learns a mapping from low-quality underwater images to visually improved reference images. The `validation` folder is used to monitor enhancement quality during training through no-reference underwater quality metrics.",
        "To test generalization beyond the training source, the implementation also uses the UIEB and RUIE benchmarks. UIEB includes 890 raw images with corresponding references and an additional challenging subset of 60 no-reference images. These two components are analytically useful because they allow both full-reference and no-reference evaluation. RUIE is used as a cross-dataset benchmark and contains multiple subsets with distinct degradation characteristics. The presence of these extra benchmarks prevents the study from depending only on in-domain performance and strengthens the external validity of the findings.",
        "The selected datasets are appropriate for a research paper because they reflect three complementary roles: supervised learning, controlled paired evaluation, and real-world no-reference testing.",
    ]
    for para in dataset_paragraphs:
        add_paragraph(doc, para)

    add_table(
        doc,
        "Table 4.1 Dataset Summary",
        ["Dataset", "Role", "Image Count", "Type", "Folder"],
        [
            ["EUVP TrainA", "Degraded training inputs", "2185", "Paired input", "data/EUVP/.../trainA"],
            ["EUVP TrainB", "Reference training targets", "2185", "Paired target", "data/EUVP/.../trainB"],
            ["EUVP Validation", "Validation during training", "130", "No-reference validation", "data/EUVP/.../validation"],
            ["UIEB raw-890", "Evaluation input set", "890", "Paired input", "data/UIEB/raw-890"],
            ["UIEB reference-890", "Evaluation reference set", "890", "Paired target", "data/UIEB/reference-890"],
            ["UIEB challenging-60", "Difficult no-reference evaluation", "60", "No-reference", "data/UIEB/challenging-60"],
            ["RUIE", "Cross-dataset evaluation", "4230", "No-reference multi-subset", "data/RUIE"],
        ],
    )

    add_picture(
        doc,
        BUNDLE / "figures" / "Figure_4_1_sample_degraded_images.png",
        Inches(5.2),
        "Figure 4.1 Sample degraded underwater images from EUVP, UIEB, and RUIE showing color cast, low contrast, and visibility loss.",
    )

    add_paragraph(
        doc,
        "From a systems perspective, the dataset design addresses two major research needs. First, the paired EUVP and UIEB subsets provide a structured environment for supervised learning and metric comparison. Second, the no-reference UIEB and RUIE images expose the model to more realistic scenes where perfect ground truth is unavailable. This is important in underwater imaging research because practical deployment rarely includes an ideal reference image for every capture."
    )

    add_heading(doc, "4.2 Training Setup")
    training_paragraphs = [
        "The training setup is implemented in PyTorch and centers on the idea of learning enhancement as a residual correction problem. During training, each degraded EUVP input image is loaded, resized, converted to RGB, transformed into a tensor, and passed through the enhancement network. The corresponding target image is used to calculate the loss. By structuring the setup in this way, the project keeps the training process consistent and computationally manageable while preserving the paired relationship required for supervised learning.",
        "The dataset loader performs several preprocessing tasks that are essential to stable training. It validates the input and target directories, filters supported image formats, matches filenames between degraded and reference images, converts images to RGB, resizes them to 256 x 256 pixels, and applies training-time augmentation. Horizontal flips, vertical flips, and color jitter on the input image help the model see moderate variations in illumination and composition. This improves robustness without changing the target supervision structure.",
        "The training loop itself follows a practical experimental design. Images are batched through a `DataLoader`, shuffled during training, and processed on a CUDA-enabled GPU when available. Validation is performed every epoch using UCIQE and UIQM, which are appropriate no-reference underwater image quality measures. A `ReduceLROnPlateau` scheduler lowers the learning rate when validation performance stops improving, helping the system move from rapid early learning to slower refinement in later epochs.",
        "The system also includes reproducibility-oriented outputs. During training it writes a CSV log, stores periodic checkpoints, saves the best-performing model according to validation UCIQE, and exports training curves for later analysis.",
    ]
    for para in training_paragraphs:
        add_paragraph(doc, para)

    add_picture(
        doc,
        BUNDLE / "figures" / "Figure_4_2_training_setup_workflow.png",
        Inches(5.2),
        "Figure 4.2 Training setup workflow showing data preparation, paired learning, validation, and checkpoint generation.",
    )

    add_paragraph(
        doc,
        "The actual training log confirms that the setup is functioning as intended. The model was trained for 100 epochs, the training loss decreased from about 0.0883 in the first epoch to about 0.0541 by the final epoch, and the best validation UCIQE was achieved early in training before later epochs stabilized around competitive values. This behavior suggests that the system quickly learns major perceptual corrections such as contrast and color balancing, then spends subsequent epochs refining those improvements."
    )

    add_heading(doc, "4.3 Hyperparameters")
    hyper_paragraphs = [
        "Hyperparameter selection has a direct impact on training stability, convergence speed, memory consumption, and output quality. In this project, the chosen values reflect a balance between computational feasibility and enhancement performance. The model uses a 256 x 256 input size so that training remains efficient on commonly available GPUs while still preserving enough visual structure for meaningful enhancement. A batch size of 8 is consistent with this design choice and reduces the risk of memory overflow during experimentation.",
        "The learning rate is set to 0.0005, which is sufficiently large for effective early optimization but not so large that the model becomes unstable. The optimizer is Adam, a widely accepted choice for image enhancement and restoration tasks because it adapts parameter updates according to gradient statistics. The training duration is fixed at 100 epochs, while checkpoints are saved every 10 epochs to support recovery and comparative analysis. These values are appropriate for a project-scale study in which reliability and interpretability are more important than brute-force experimentation.",
        "An important part of the hyperparameter design is the combined loss function. The training script assigns weights of 0.6 to L1 reconstruction loss, 0.3 to edge loss, and 0.1 to color loss. This weighting scheme reveals the intended behavior of the system. L1 loss ensures that outputs remain close to the target image in a broad pixel-wise sense. Edge loss encourages structural sharpness, which is useful for recovering object boundaries and local detail in hazy underwater scenes. Color loss penalizes undesirable channel imbalance and therefore supports the restoration of more natural visual appearance.",
    ]
    for para in hyper_paragraphs:
        add_paragraph(doc, para)

    add_table(
        doc,
        "Table 4.2 Hyperparameter Summary",
        ["Parameter", "Configured Value", "Rationale"],
        [
            ["Image size", "256 x 256", "Balances detail retention and computational cost"],
            ["Batch size", "8", "Fits practical GPU memory while keeping updates stable"],
            ["Epochs", "100", "Allows convergence monitoring across a full training cycle"],
            ["Learning rate", "0.0005", "Supports steady optimization with Adam"],
            ["Save every", "10 epochs", "Preserves intermediate recovery points"],
            ["Loss weights", "L1 0.6, Edge 0.3, Color 0.1", "Combines fidelity, sharpness, and color correction"],
            ["Architecture", "CNN encoder-decoder + MobileViT bottleneck", "Captures local and global underwater distortions"],
        ],
    )

    add_paragraph(
        doc,
        "From a system analysis viewpoint, these hyperparameters do not appear arbitrary. They are closely aligned with the repository design in `config.py` and `training/train.py`, and they support the specific challenge of underwater enhancement, where global color restoration and local detail recovery must be learned at the same time."
    )

    add_heading(doc, "4.4 Software Requirements")
    software_paragraphs = [
        "The software environment of the proposed system is intentionally lightweight and modular. Python is used as the implementation language because of its maturity in machine learning research and the availability of scientific libraries. PyTorch provides the deep learning framework for model definition, automatic differentiation, optimization, and GPU execution. TorchVision supports image transformations, while Pillow and OpenCV handle image I/O and low-level processing operations. NumPy and scikit-image contribute numerical and image-quality utilities, and Matplotlib is used for training curves and result visualization.",
        "The project can be executed in a standard development setup such as VS Code or Jupyter Notebook, which is helpful for both experimentation and classroom-based project supervision. The repository also uses a single command entry point through `main.py`, simplifying training, inference, and evaluation. This design is valuable in a research setting because it reduces command complexity, lowers the chance of user error, and improves the reproducibility of experimental steps.",
    ]
    for para in software_paragraphs:
        add_paragraph(doc, para)

    add_table(
        doc,
        "Table 4.3 Software Requirements",
        ["Component", "Requirement"],
        [
            ["Operating system", "Windows or Linux"],
            ["Programming language", "Python"],
            ["Deep learning framework", "PyTorch"],
            ["Image libraries", "OpenCV, Pillow, TorchVision"],
            ["Numerical libraries", "NumPy, scikit-image"],
            ["Visualization", "Matplotlib"],
            ["Development tools", "VS Code or Jupyter Notebook"],
            ["Dataset formats", "JPG, JPEG, PNG, BMP, TIF, TIFF"],
            ["Acceleration", "CUDA when available"],
        ],
    )

    add_paragraph(
        doc,
        "These requirements keep the implementation accessible while still supporting GPU acceleration and publication-oriented figure generation."
    )

    add_heading(doc, "4.5 Implementation Process")
    implementation_paragraphs = [
        "The implementation process begins with repository-level configuration. Dataset paths, output paths, and training hyperparameters are centralized in `config.py`, which reduces duplication and makes the system easier to maintain. After configuration, the user can run training, inference, evaluation, or environment checking through `main.py`. This command-routing design functions as a simple control layer for the entire project and ensures that major workflows are launched from one consistent interface.",
        "The next stage is data loading and preprocessing. The custom dataset class reads supported image files, aligns paired filenames, resizes all images to a fixed spatial size, converts them to tensors, and optionally performs augmentation. Once loaded, the images are passed to the enhancement model defined in `models/cnn_model.py`. The architecture uses two convolutional encoder stages, a MobileViT-style bottleneck for broader contextual reasoning, and two decoder stages with skip connections. The final output is generated through a residual formulation in which the network predicts a correction that is added to the input image before a sigmoid activation is applied.",
        "After forward propagation, the training script computes the combined loss, performs backpropagation, clips gradients, updates parameters through Adam, and then validates the current model using UCIQE and UIQM. The system saves both the best checkpoint and the final checkpoint, while also preserving periodic epoch snapshots. When training is complete, the same repository supports inference on single images, batches, and validation folders, as well as evaluation on EUVP, UIEB, and RUIE benchmarks.",
        "The final part of the implementation process is result generation for analysis. Quantitative summaries are exported to text and CSV files, and qualitative comparisons are stored as image outputs. This combination is academically meaningful because underwater enhancement cannot be judged by numbers alone. A visually improved image may not always maximize full-reference metrics, while a metric-friendly output may still look unnatural to a human observer. The implementation therefore supports both perspectives and encourages balanced interpretation.",
    ]
    for para in implementation_paragraphs:
        add_paragraph(doc, para)

    add_picture(
        doc,
        BUNDLE / "figures" / "Figure_4_5_system_implementation_workflow.png",
        Inches(5.2),
        "Figure 4.3 Overall implementation process of the proposed underwater image enhancement system.",
    )

    add_heading(doc, "4.6 Analytical Discussion of System Behavior")
    discussion_paragraphs = [
        "The implemented system shows a clear orientation toward perceptual underwater enhancement. On EUVP validation, UCIQE improved from 4.8264 to 5.7364 and UIQM improved from 0.5900 to 0.6126. On UIEB paired data, the model improved UCIQE from 4.3888 to 5.1333 and slightly improved UIQM from 0.5082 to 0.5085, although PSNR and SSIM decreased modestly. On the UIEB challenging subset, UCIQE increased from 3.1536 to 3.9637 and UIQM increased from 0.4229 to 0.4400. RUIE overall evaluation also showed strong UCIQE improvement from 2.1926 to 3.4089. These values indicate that the system consistently enhances underwater perceptual quality across datasets.",
        "This pattern is analytically important. In underwater image enhancement, exact agreement with a reference image does not always coincide with improved visual realism. A slight reduction in PSNR or SSIM can occur when the model introduces stronger color correction or contrast enhancement that better matches human perception but differs from the specific paired target.",
        "The system design helps explain this behavior. The residual architecture allows the network to focus on corrective information rather than reconstructing the entire image from scratch, while the edge-aware and color-aware loss terms push the model toward sharper structure and better color balance.",
        "Some limitations are also visible in the analysis. The fixed 256 x 256 input size may suppress fine details from high-resolution scenes, and the training source remains primarily EUVP even though cross-dataset evaluation is included. Nevertheless, the overall system is well structured, technically coherent, and supported by measurable outputs, which makes it suitable for inclusion in a research paper."
    ]
    for para in discussion_paragraphs:
        add_paragraph(doc, para)

    add_heading(doc, "4.7 Summary")
    summary_paragraphs = [
        "This chapter presented the system analysis of the proposed underwater image enhancement framework by examining the dataset base, the training setup, the chosen hyperparameters, the software environment, and the implementation workflow. The analysis showed that the project is not only a model definition but a complete experimental pipeline built around paired training, controlled validation, and cross-dataset evaluation.",
        "The chapter also demonstrated that the implementation choices are consistent with the research objective. Benchmark datasets provide credible experimental inputs, the preprocessing pipeline standardizes images for supervised learning, the hyperparameters support stable optimization, and the software stack enables repeatable execution. Finally, the exported metrics, checkpoints, curves, and qualitative results make the system suitable for scholarly presentation and later comparison with alternative methods.",
        "Overall, the proposed Section 4 establishes a strong technical foundation for the subsequent results and conclusion chapters. It explains how the underwater image enhancement system is constructed, why the chosen design is academically defensible, and how the implementation supports both engineering reliability and research-style analysis."
    ]
    for para in summary_paragraphs:
        add_paragraph(doc, para)

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
